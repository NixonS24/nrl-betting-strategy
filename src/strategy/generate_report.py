"""
NRL Bias Research — Report Generator
Produces a formatted Word document (.docx) with charts embedded.

Usage:
    python src/strategy/generate_report.py
Output:
    data/processed/NRL_Bias_Research_Report.docx
"""

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import io
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── import strategy module ────────────────────────────────────────────────
import sys
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))
from src.strategy.venue_bias import (
    load_data, compute_venue_baselines, run_backtest, summarise,
    BACK_HOME_VENUES, FADE_HOME_VENUES, MIN_EDGE, MIN_ODDS, MAX_ODDS, FLAT_STAKE
)
from src.analysis.bias_analysis import *   # noqa — runs the analysis, populates globals

ROOT   = Path(__file__).resolve().parents[2]
OUT    = ROOT / "data" / "processed" / "NRL_Bias_Research_Report.docx"
CHARTS = ROOT / "data" / "processed" / "charts"
CHARTS.mkdir(exist_ok=True)

BRAND_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
BRAND_GREEN = RGBColor(0x37, 0x86, 0x44)
BRAND_RED   = RGBColor(0xC0, 0x39, 0x2B)
BRAND_GREY  = RGBColor(0x60, 0x60, 0x60)


# ── Helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def add_heading(doc, text, level=1, colour=BRAND_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = colour
    return h


def add_table(doc, headers, rows, col_widths=None, header_bg="1F497D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_bg)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
            if ri % 2 == 0:
                set_cell_bg(cells[ci], "EBF0F7")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def fig_to_docx(fig, doc, width=6.0):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    plt.close(fig)


# ── Charts ────────────────────────────────────────────────────────────────

def chart_venue_home_win_rate(df):
    sub = df.copy()
    sub["home_win"] = (sub["result"] == "home_win").astype(int)
    by_v = sub.groupby("venue_name")["home_win"].agg(["mean","count"])
    by_v = by_v[by_v["count"] >= 30].sort_values("mean", ascending=True)
    overall = sub["home_win"].mean()

    fig, ax = plt.subplots(figsize=(9, 6))
    colours = ["#C0392B" if v in FADE_HOME_VENUES else
               "#1F497D" if v in BACK_HOME_VENUES else
               "#95A5A6"
               for v in by_v.index]
    bars = ax.barh(by_v.index, by_v["mean"] * 100, color=colours, edgecolor="white")
    ax.axvline(overall * 100, color="#E67E22", linewidth=2, linestyle="--",
               label=f"League avg ({overall*100:.1f}%)")
    ax.set_xlabel("Home Win Rate (%)", fontsize=11)
    ax.set_title("Home Win Rate by Venue (≥30 matches, 1998–2025)", fontsize=12, fontweight="bold")
    ax.legend(fontsize=9)

    from matplotlib.patches import Patch
    legend_handles = [
        Patch(color="#1F497D", label="Target: Back Home"),
        Patch(color="#C0392B", label="Target: Fade Home"),
        Patch(color="#95A5A6", label="No signal"),
        plt.Line2D([0],[0], color="#E67E22", linewidth=2, linestyle="--",
                   label=f"League avg ({overall*100:.1f}%)"),
    ]
    ax.legend(handles=legend_handles, fontsize=9, loc="lower right")
    ax.xaxis.set_major_formatter(mticker.FormatStrFormatter("%.0f%%"))
    fig.tight_layout()
    return fig


def chart_equity_curve(bets: pd.DataFrame, label: str):
    if bets.empty:
        return None
    fig, ax = plt.subplots(figsize=(9, 4))
    ax.plot(bets["date"], bets["bankroll"], color="#1F497D", linewidth=2)
    ax.axhline(10_000, color="#95A5A6", linewidth=1, linestyle="--", label="Starting bankroll")
    ax.fill_between(bets["date"], 10_000, bets["bankroll"],
                    where=bets["bankroll"] >= 10_000, alpha=0.15, color="#1F497D")
    ax.fill_between(bets["date"], 10_000, bets["bankroll"],
                    where=bets["bankroll"] < 10_000, alpha=0.15, color="#C0392B")
    ax.set_ylabel("Bankroll ($)", fontsize=11)
    ax.set_title(f"Equity Curve — {label} (flat ${FLAT_STAKE} stake, $10k start)", fontsize=12, fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"${x:,.0f}"))
    ax.legend(fontsize=9)
    fig.tight_layout()
    return fig


def chart_season_pnl(bets: pd.DataFrame, label: str):
    if bets.empty or "season" not in bets.columns:
        return None
    by_s = bets.groupby("season")["profit"].sum().reset_index()
    fig, ax = plt.subplots(figsize=(9, 4))
    colours = ["#378644" if p >= 0 else "#C0392B" for p in by_s["profit"]]
    ax.bar(by_s["season"].astype(str), by_s["profit"], color=colours, edgecolor="white")
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_ylabel("Profit / Loss ($)", fontsize=11)
    ax.set_title(f"Annual P&L — {label}", fontsize=12, fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"${x:,.0f}"))
    plt.xticks(rotation=45, ha="right")
    fig.tight_layout()
    return fig


def chart_form_buckets(form_df):
    buckets = ["strong_away","mild_away","neutral","mild_home","strong_home"]
    labels  = ["Strong\nAway Form","Mild\nAway Form","Neutral","Mild\nHome Form","Strong\nHome Form"]
    rates = []
    for b in buckets:
        sub = form_df[form_df["form_advantage"] == b]
        rates.append(sub["home_win_bin"].mean() * 100 if len(sub) else 0)

    overall_hw = form_df["home_win_bin"].mean() * 100
    fig, ax = plt.subplots(figsize=(8, 4))
    colours = ["#C0392B","#E07050","#95A5A6","#5B9BD5","#1F497D"]
    ax.bar(labels, rates, color=colours, edgecolor="white")
    ax.axhline(overall_hw, color="#E67E22", linewidth=2, linestyle="--",
               label=f"League avg ({overall_hw:.1f}%)")
    ax.set_ylabel("Home Win Rate (%)", fontsize=11)
    ax.set_title("Home Win Rate by Recent Form Advantage (last 5 games)", fontsize=12, fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FormatStrFormatter("%.0f%%"))
    ax.legend(fontsize=9)
    fig.tight_layout()
    return fig


# ── Build document ────────────────────────────────────────────────────────

def build_report():
    # ── Load data & run backtests ─────────────────────────────────────────
    df = load_data()
    baselines = compute_venue_baselines(df)
    bk_bets = run_backtest(df, baselines, "bk_home_close", "bk_away_close", "Bookmaker")
    bf_bets = run_backtest(df, baselines, "bf_home_open",  "bf_away_open",  "Betfair")
    bk = summarise(bk_bets)
    bf = summarise(bf_bets)

    # Form analysis (recompute for charts)
    form_df = df[df["home_form_last5"].notna() & df["away_form_last5"].notna()].copy()
    form_df["form_diff"] = form_df["home_form_last5"] - form_df["away_form_last5"]
    form_df["home_win_bin"] = (form_df["result"] == "home_win").astype(int)
    form_df["form_advantage"] = pd.cut(
        form_df["form_diff"],
        bins=[-1.1,-0.4,-0.1,0.1,0.4,1.1],
        labels=["strong_away","mild_away","neutral","mild_home","strong_home"]
    )

    # Betfair venue edge table
    bf_venue_df = df[df["bf_home_open"].notna() & df["venue_name"].notna()].copy()
    bf_venue_df["home_win_bin"] = (bf_venue_df["result"] == "home_win").astype(int)
    bf_venue_df["implied_home"] = 1 / bf_venue_df["bf_home_open"]
    bf_by_venue = (
        bf_venue_df.groupby("venue_name")
        .agg(n=("home_win_bin","count"),
             actual_hw=("home_win_bin","mean"),
             implied_hw=("implied_home","mean"))
        .query("n >= 10")
        .assign(edge=lambda x: x["actual_hw"] - x["implied_hw"])
        .sort_values("actual_hw", ascending=False)
        .reset_index()
    )

    # ── Document ──────────────────────────────────────────────────────────
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Cover ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("NRL Gambling Market Bias Research")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BRAND_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Statistical Analysis & Recommended Betting Strategy")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = BRAND_GREY

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Generated: {date.today().strftime('%d %B %Y')}").font.color.rgb = BRAND_GREY

    doc.add_paragraph()
    doc.add_paragraph()

    # Executive summary box
    exec_p = doc.add_paragraph()
    exec_run = exec_p.add_run(
        "EXECUTIVE SUMMARY\n\n"
        "This report identifies a statistically significant and financially exploitable "
        "Venue/Home Bias in the NRL betting market. Analysis of 5,435 matches (1998–2025) "
        "shows that specific venues — most notably AAMI Park (Melbourne Storm's home ground) — "
        "produce home win rates materially higher than market-implied probabilities, generating "
        f"a consistent edge of +10–12%. A backtested flat-stake strategy across "
        f"{bk.get('n', 0)} qualifying bookmaker bets (2009–2025) returned an ROI of "
        f"{bk.get('roi', 0):.1%}. Draw Bias was not significant. Form Bias is real but "
        "appears to be priced in by the market."
    )
    exec_run.font.size = Pt(10)
    exec_run.font.italic = True

    doc.add_page_break()

    # ── 1. Introduction ───────────────────────────────────────────────────
    add_heading(doc, "1. Introduction", level=1)
    doc.add_paragraph(
        "This research project investigates three hypothesised inefficiencies in the "
        "Australian NRL (National Rugby League) betting market:\n\n"
        "  1.  Draw Bias — the market underprices the probability of drawn matches "
        "between high-profile clubs.\n"
        "  2.  Form/Momentum Bias — the market over- or under-reacts to recent team "
        "performance, creating exploitable mispricings.\n"
        "  3.  Venue/Home Bias — the market systematically miscalculates home advantage "
        "at specific venues."
    )
    doc.add_paragraph(
        "Data sources used:"
    )
    add_table(doc,
        headers=["Source", "Coverage", "Odds?", "Rows"],
        rows=[
            ["AusSportsBetting (nrl.xlsx)", "2009–2025", "Yes — H2H, line, totals", "3,469"],
            ["uselessnrlstats (GitHub CSVs)", "1998–2025", "No (match results only)", "5,435 NRL"],
            ["Betfair Automation Hub CSVs", "2021–2025", "Yes — exchange prices", "~2,000"],
        ],
        col_widths=[2.5, 1.4, 2.0, 0.8]
    )

    # ── 2. Draw Bias ──────────────────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "2. Draw Bias", level=1)
    add_heading(doc, "Finding: Not Statistically Significant (p = 0.62)", level=2,
                colour=BRAND_RED)

    doc.add_paragraph(
        "NRL matches end in a draw extremely rarely — only 55 times across 5,435 matches "
        "(1.01%). There is no meaningful difference in draw frequency between top-8 club "
        "matchups (1.17%) and other games (0.96%), confirmed by a chi-square test "
        "(χ² = 0.25, p = 0.62).\n\n"
        "Additionally, Betfair does not operate a standalone draw market for NRL head-to-head "
        "betting. The sum of implied home and away probabilities averages 1.008, leaving a "
        "residual of −0.01 — consistent with bookmaker overround, not a priced draw outcome.\n\n"
        "Conclusion: Draw Bias is not exploitable in the NRL market."
    )

    # ── 3. Form / Momentum Bias ───────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "3. Form / Momentum Bias", level=1)
    add_heading(doc, "Finding: Statistically Significant but Market-Adjusted (p < 0.0001)", level=2,
                colour=BRAND_GREEN)

    doc.add_paragraph(
        "Recent form is a strong predictor of match outcome: teams with a superior record "
        "over the last five games win significantly more often, and this effect is large "
        "enough to be statistically unambiguous."
    )
    add_table(doc,
        headers=["Form Situation", "Home Win Rate", "Matches"],
        rows=[
            ["Strong away form advantage", "42.4%", "839"],
            ["Mild away form advantage",   "49.8%", "1,332"],
            ["Neutral",                    "59.8%", "1,196"],
            ["Mild home form advantage",   "65.6%", "1,630"],
            ["Strong home form advantage", "73.2%", "422"],
        ],
        col_widths=[2.8, 1.5, 1.0]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "However, the Betfair market already accounts for this. The correlation between "
        "recent form differential and Betfair's implied home-win probability is r = 0.66 "
        "(p < 0.0001) — the market has largely priced in form, leaving limited residual edge.\n\n"
        "Conclusion: Form Bias is real but the market tracks it closely. Monitoring for "
        "situations where form changes rapidly between market open and close may yield "
        "small edges, but this is not the primary opportunity."
    )

    fig_form = chart_form_buckets(form_df)
    fig_to_docx(fig_form, doc, width=5.8)

    # ── 4. Venue / Home Bias ─────────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "4. Venue / Home Bias", level=1)
    add_heading(doc, "Finding: Highly Significant — Primary Opportunity (p < 0.0001)", level=2,
                colour=BRAND_GREEN)

    doc.add_paragraph(
        "Home win rates vary dramatically and significantly across NRL venues "
        "(one-way ANOVA: F = 4.33, p < 0.0001). The league-wide home win rate is 57.4%, "
        "but several venues diverge substantially and persistently from this baseline."
    )

    venue_rows = []
    overall_hw = df[df["venue_name"].notna()]["home_win"].mean()
    by_venue = (
        df[df["venue_name"].notna()]
        .assign(home_win=lambda x: (x["result"]=="home_win").astype(int))
        .groupby("venue_name")["home_win"]
        .agg(["mean","count"])
    )
    by_venue = by_venue[by_venue["count"] >= 30].sort_values("mean", ascending=False)
    for v, row in by_venue.iterrows():
        direction = "↑ Above" if row["mean"] > overall_hw else "↓ Below"
        venue_rows.append([v, f"{row['mean']*100:.1f}%", int(row["count"]), direction])

    add_table(doc,
        headers=["Venue", "Home Win Rate", "Matches", "vs League Avg"],
        rows=venue_rows,
        col_widths=[2.8, 1.3, 0.9, 1.0]
    )
    doc.add_paragraph()
    fig_venue = chart_venue_home_win_rate(df)
    fig_to_docx(fig_venue, doc, width=6.2)

    doc.add_paragraph()
    add_heading(doc, "4.1 Market Mispricing at Key Venues", level=2)
    doc.add_paragraph(
        "Cross-referencing Betfair exchange opening prices with actual outcomes reveals "
        "that the market consistently underestimates Melbourne Storm's home advantage at "
        "AAMI Park — the largest and most consistent edge found in this study."
    )

    bf_rows = []
    for _, row in bf_by_venue.iterrows():
        edge_str = f"+{row['edge']*100:.1f}%" if row["edge"] > 0 else f"{row['edge']*100:.1f}%"
        bf_rows.append([row["venue_name"], int(row["n"]),
                        f"{row['actual_hw']*100:.1f}%",
                        f"{row['implied_hw']*100:.1f}%",
                        edge_str])

    add_table(doc,
        headers=["Venue", "Matches", "Actual Home Win%", "Market Implied%", "Edge"],
        rows=bf_rows,
        col_widths=[2.5, 0.8, 1.4, 1.4, 0.9]
    )

    # ── 5. Recommended Strategy ───────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Recommended Strategy: Venue Bias Exploitation", level=1)

    doc.add_paragraph(
        "Based on the statistical analysis, we recommend a selective betting strategy "
        "targeting confirmed venue mispricings. The strategy has two legs:"
    )

    add_heading(doc, "Leg 1 — Back Home Team", level=2, colour=BRAND_GREEN)
    doc.add_paragraph(
        "Back the home team at the following venues when the bookmaker's closing implied "
        f"probability is more than {MIN_EDGE*100:.0f} percentage points below the venue's "
        "historical home win rate:"
    )
    back_rows = []
    for v in sorted(BACK_HOME_VENUES):
        row = baselines[baselines["venue_name"] == v]
        if not row.empty:
            back_rows.append([v, f"{row['base_hw_rate'].values[0]*100:.1f}%",
                               int(row["n_matches"].values[0]),
                               "Backing home team"])
    add_table(doc,
        headers=["Venue", "Historical Home Win%", "Matches", "Action"],
        rows=back_rows,
        col_widths=[2.5, 1.8, 1.0, 1.7]
    )

    add_heading(doc, "Leg 2 — Fade Home Team (Back Away)", level=2, colour=BRAND_RED)
    doc.add_paragraph(
        "Back the away team at the following venues where the home team historically "
        "underperforms market expectations:"
    )
    fade_rows = []
    for v in sorted(FADE_HOME_VENUES):
        row = baselines[baselines["venue_name"] == v]
        if not row.empty:
            fade_rows.append([v, f"{row['base_hw_rate'].values[0]*100:.1f}%",
                               int(row["n_matches"].values[0]),
                               "Backing away team"])
    add_table(doc,
        headers=["Venue", "Historical Home Win%", "Matches", "Action"],
        rows=fade_rows,
        col_widths=[2.5, 1.8, 1.0, 1.7]
    )

    add_heading(doc, "Bet Selection Rules", level=2)
    add_table(doc,
        headers=["Parameter", "Value", "Rationale"],
        rows=[
            ["Minimum edge", f"{MIN_EDGE*100:.0f}%", "Implied prob must be this far below historical rate"],
            ["Minimum odds",  f"${MIN_ODDS:.2f}",   "Avoid heavily-priced favourites — high variance"],
            ["Maximum odds",  f"${MAX_ODDS:.2f}",   "Limit tail risk on longshots"],
            ["Staking",       "Flat or ¼ Kelly",    "Flat stake for simplicity; Kelly for growth"],
            ["Market",        "Bookmaker closing / Betfair pre-match", "Closing odds are most efficient"],
            ["Trigger",       "Venue + edge condition met",   "Do not bet every home game — only when edge exists"],
        ],
        col_widths=[1.8, 1.5, 3.2]
    )

    # ── 6. Backtest Results ───────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "6. Backtest Results", level=1)

    add_heading(doc, f"6.1 Bookmaker Odds — Primary Backtest (2009–2025)", level=2)
    if bk:
        add_table(doc,
            headers=["Metric", "Value"],
            rows=[
                ["Total bets placed",    str(bk["n"])],
                ["Winning bets",         f"{bk['wins']} ({bk['win_rate']:.1%})"],
                ["Total staked",         f"${bk['staked']:,.0f}"],
                ["Total profit",         f"${bk['profit']:,.2f}"],
                ["Return on Investment", f"{bk['roi']:.2%}"],
                ["Max drawdown",         f"${bk['max_drawdown']:,.2f}"],
                ["Average odds",         f"{bk['avg_odds']:.3f}"],
                ["Average edge",         f"{bk['avg_edge']:.2%}"],
            ],
            col_widths=[2.5, 2.5]
        )
        doc.add_paragraph()

        fig_eq = chart_equity_curve(bk_bets, "Bookmaker (2009–2025)")
        if fig_eq:
            fig_to_docx(fig_eq, doc, width=6.0)

        doc.add_paragraph()
        fig_pnl = chart_season_pnl(bk_bets, "Bookmaker")
        if fig_pnl:
            fig_to_docx(fig_pnl, doc, width=6.0)

        doc.add_paragraph()
        add_heading(doc, "Results by Venue", level=3)
        v_rows = []
        for v, row in bk["by_venue"].iterrows():
            v_rows.append([v, int(row["n"]), int(row["wins"]),
                           f"${row['profit']:,.2f}", f"{row['roi']:.1%}"])
        add_table(doc,
            headers=["Venue", "Bets", "Wins", "Profit", "ROI"],
            rows=v_rows,
            col_widths=[2.8, 0.7, 0.7, 1.1, 0.9]
        )

        doc.add_paragraph()
        add_heading(doc, "All Individual Bets", level=3)
        bet_rows = []
        for _, b in bk_bets.iterrows():
            bet_rows.append([
                b["date"].strftime("%Y-%m-%d"),
                b["home_team"], b["away_team"],
                b["venue"], b["bet_type"].replace("_", " ").title(),
                f"${b['odds']:.2f}", f"{b['edge']*100:.1f}%",
                b["result"].replace("_", " ").title(),
                "✓ Win" if b["won"] else "✗ Loss",
                f"${b['profit']:+.2f}",
                f"${b['cumulative_profit']:+.2f}",
            ])
        add_table(doc,
            headers=["Date","Home","Away","Venue","Bet","Odds","Edge","Result","W/L","Profit","Cumulative"],
            rows=bet_rows,
            col_widths=[0.9, 1.3, 1.3, 1.5, 0.9, 0.6, 0.55, 0.9, 0.65, 0.65, 0.85]
        )

    add_heading(doc, "6.2 Betfair Exchange — Validation (2021–2025)", level=2)
    if bf:
        add_table(doc,
            headers=["Metric", "Value"],
            rows=[
                ["Total bets",    str(bf["n"])],
                ["Win rate",      f"{bf['win_rate']:.1%}"],
                ["Total profit",  f"${bf['profit']:,.2f}"],
                ["ROI",           f"{bf['roi']:.2%}"],
                ["Max drawdown",  f"${bf['max_drawdown']:,.2f}"],
            ],
            col_widths=[2.5, 2.5]
        )
        doc.add_paragraph()
        fig_bf = chart_equity_curve(bf_bets, "Betfair Exchange (2021–2025)")
        if fig_bf:
            fig_to_docx(fig_bf, doc, width=6.0)

    # ── 7. Limitations ────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7. Limitations & Risk Factors", level=1)
    add_table(doc,
        headers=["Risk", "Description", "Mitigation"],
        rows=[
            ["Small sample",
             "Only ~10–50 qualifying bets per venue per decade. "
             "Betfair validation sample is particularly small (10 bets).",
             "Accumulate more data over 2025–2026 season before scaling stakes."],
            ["Look-ahead bias",
             "Venue base rates use full historical data; real-time application "
             "requires rolling baselines.",
             "Re-run pipeline at season start with prior seasons only."],
            ["Market efficiency drift",
             "Edges tend to shrink as they become known. "
             "AAMI Park is a well-known Melbourne Storm fortress.",
             "Monitor edge size quarterly; exit if implied prob converges to baseline."],
            ["Team relocation / venue change",
             "AAMI Park edge is tied to Melbourne Storm. If they move, "
             "the historical baseline no longer applies.",
             "Track venue assignments per team each season."],
            ["Bookmaker limits",
             "Winning punters are restricted by bookmakers.",
             "Use Betfair exchange (no limits) or multiple bookmaker accounts."],
            ["Selection bias",
             "Games at target venues are not random — they reflect "
             "home-team strength.",
             "The market already adjusts for team quality; the residual edge "
             "is the venue effect net of team quality."],
        ],
        col_widths=[1.5, 2.8, 2.2]
    )

    # ── 8. Recommended Next Steps ─────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "8. Recommended Next Steps", level=1)
    add_table(doc,
        headers=["Priority", "Action", "Owner"],
        rows=[
            ["1 — High",   "Apply rolling baseline (prior seasons only) to eliminate look-ahead bias", "Data Engineer"],
            ["2 — High",   "Extend Betfair backtest using All_Markets CSVs for additional sample size",   "Data Engineer"],
            ["3 — Medium", "Monitor 2025 NRL season in real time — log qualifying bets and outcomes",      "Strategy"],
            ["4 — Medium", "Build OddsPortal scraper to fill 2009–2012 bookmaker odds gap",               "Data Sourcer"],
            ["4 — Medium", "Test Form Bias as a filter on top of Venue Bias (form + venue combined)",     "Data Scientist"],
            ["5 — Low",    "Investigate Campbelltown Sports Stadium fade further — extreme signal",        "Data Scientist"],
            ["5 — Low",    "Create Jupyter EDA notebook with interactive venue maps",                     "Data Scientist"],
        ],
        col_widths=[1.1, 4.0, 1.4]
    )

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(OUT)
    print(f"Report written → {OUT}")
    return OUT


if __name__ == "__main__":
    # Suppress output from imported analysis module
    import contextlib, os
    with open(os.devnull, "w") as devnull:
        with contextlib.redirect_stdout(devnull):
            build_report()
    print(f"\nReport saved to: {OUT}")
